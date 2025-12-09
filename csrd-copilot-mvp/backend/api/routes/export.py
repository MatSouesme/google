from fastapi import APIRouter, HTTPException, Response
from pydantic import BaseModel
from fpdf import FPDF
from docx import Document
import markdown
import io
import re

router = APIRouter()

class ExportRequest(BaseModel):
    content: str
    title: str = "CSRD Report"

def clean_markdown(text):
    """Removes markdown formatting for plain text usage."""
    # Remove bold/italic
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove headers
    text = re.sub(r'#+\s', '', text)
    return text

@router.post("/export/pdf")
def export_pdf(request: ExportRequest):
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Title
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, request.title, ln=True, align="C")
        pdf.ln(10)
        
        # Content
        # FPDF doesn't support Markdown natively well, so we do basic cleaning
        # For a real MVP, we might want to use a library that converts HTML to PDF like weasyprint,
        # but that has heavy dependencies. FPDF is safe.
        
        pdf.set_font("Arial", size=12)
        
        # Split by lines to handle basic formatting
        lines = request.content.split('\n')
        for line in lines:
            # Headers
            if line.startswith('# '):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, line.replace('# ', ''), ln=True)
                pdf.set_font("Arial", size=12)
            elif line.startswith('## '):
                pdf.set_font("Arial", "B", 13)
                pdf.cell(0, 10, line.replace('## ', ''), ln=True)
                pdf.set_font("Arial", size=12)
            elif line.startswith('### '):
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, line.replace('### ', ''), ln=True)
                pdf.set_font("Arial", size=12)
            else:
                # Regular text
                # Handle bolding (simple removal for now as FPDF mix is hard)
                clean_line = clean_markdown(line)
                # Multi_cell for text wrapping
                pdf.multi_cell(0, 10, clean_line)
                
        
        output = io.BytesIO()
        pdf.output(output)
        output.seek(0)
        
        return Response(
            content=output.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={request.title}.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

@router.post("/export/docx")
def export_docx(request: ExportRequest):
    try:
        doc = Document()
        doc.add_heading(request.title, 0)
        
        # Simple Markdown parsing
        lines = request.content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            else:
                # Remove markdown bold/italic for docx (or implement rich text parsing)
                clean_line = clean_markdown(line)
                if clean_line.strip():
                    doc.add_paragraph(clean_line)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={request.title}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")
